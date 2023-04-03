import streamlit as st
from utils import defaultConfig,process_image
from streamlit_webrtc import webrtc_streamer, VideoTransformerBase
import threading
from typing import Union
import av
from model import model, tableModel
from PIL import Image
import cv2
import numpy as np
import uuid
from streamlit_option_menu import option_menu
from io import BytesIO
from docx import Document

defaultConfig(1)
st.image("./media/logo.png")
st.text("To convert a Math Equation to its equivalent LaTeX Code, upload image of the equation from your device")
IMAGENAME = ''

class VideoTransformer(VideoTransformerBase):
    frame_lock: threading.Lock  

    out_image: Union[np.ndarray, None]

    def __init__(self) -> None:
        self.frame_lock = threading.Lock()
        self.out_image = None

    def transform(self, frame: av.VideoFrame) -> np.ndarray:
        out_image = frame.to_ndarray(format="bgr24")

        with self.frame_lock:
            self.out_image = out_image

        height, width, _ = out_image.shape
        x = int(width/4)
        y = int(height/4)
        w = int(width/2)
        h = int(height/2)
        
        out_image = cv2.rectangle(out_image, (x, y), (x+w, y+h), (0, 255, 0), 2)
        return out_image
    
def app():

    selected = option_menu(
        None, 
        options=["Scan","Img","Table", "FAQ","Help","About"],
        icons=["webcam", "file-image", "table", "patch-question", "info-circle", "file-earmark-person"],
        menu_icon="cast",
        default_index=1,
        orientation="horizontal")
    
    if selected == "Scan":
        webcam = st.empty()

        with webcam:
            wrtc = webrtc_streamer(key="snapshot", video_transformer_factory=VideoTransformer)
    
        if wrtc.video_transformer:
            snap = st.button("Snapshot")

            if snap:
                with wrtc.video_transformer.frame_lock:
                    out_image = wrtc.video_transformer.out_image

                if out_image is not None:
                    height, width, _ = out_image.shape
                    x = int(width/4)
                    y = int(height/4)
                    w = int(width/2)
                    h = int(height/2)
                    
                    cropped_image = out_image[y+2:y+h-2, x+2:x+w-2]
                    rgbImage = cv2.cvtColor(cropped_image, cv2.COLOR_BGR2RGB)
                    image = Image.fromarray(rgbImage)

                    uid = uuid.uuid4().hex
                    filename = f'./images/{uid}.jpg'

                    image.save(filename)
                    global IMAGENAME
                    IMAGENAME = filename
                    webcam.empty()
                        
                else:
                    st.warning("No frames available yet.")
        
        if len(IMAGENAME) > 0:
            st.text("Scaned Image")
            image = Image.open(IMAGENAME)
            display_image = np.array(image)
            st.image(display_image, channels="BGR")

            if image is not None:
                generatedLatex = model(image)
                st.balloons()
                st.text("Generated Latex")
                st.code(generatedLatex, language='latex')
                st.text("Complied Latex")
                st.latex(generatedLatex)

                document = Document()
                document.add_paragraph(generatedLatex)
                stream = BytesIO()

                document.save(stream)
                stream.seek(0)
                
                st.download_button(
                    label="Download LaTeX as Word file",
                    data= stream,
                    file_name='generated_latex.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

    if selected == "Img":
        imageHead = Image.open('./images/img-latex.jpg')

        st.image(imageHead)
        
        imageUploaded = st.file_uploader('Choose an image file', type=['jpg','png', 'jpeg'])
        if imageUploaded is not None:
            process_image(imageUploaded)
            arrayImage = Image.open(imageUploaded)
            display_image = np.array(arrayImage)

            generatedLatex = model(arrayImage)
            st.balloons()
            st.write("Generated Latex")
            st.code(generatedLatex, language='latex')
            st.write("Complied Latex")
            st.latex(generatedLatex)
            
            document = Document()
            document.add_paragraph(generatedLatex)
            stream = BytesIO()

            document.save(stream)
            stream.seek(0)
            
            st.download_button(
                label="Download LaTeX as Word file",
                data= stream,
                file_name='generated_latex.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    
    if selected == "Table":
        imageHead = Image.open('./images/tabel-latex.jpg')

        st.image(imageHead)

        pdfUploaded = st.file_uploader('Choose an PDF file', type=['pdf'])
        if pdfUploaded:
            re = tableModel(pdfUploaded)
            st.balloons()
            st.write("Generated Latex")
            generatedTableLatex = re[0].to_latex()
            generatedTabledf = re[0]
            st.code(generatedTableLatex, language="latex")
            print(generatedTableLatex)
            st.write("Complied Latex")
            st.dataframe(generatedTabledf)
            
            document = Document()
            document.add_paragraph(generatedTableLatex)
            stream = BytesIO()

            document.save(stream)
            stream.seek(0)
            
            st.download_button(
                label="Download LaTeX as Word file",
                data= stream,
                file_name='generated_latex.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    if selected == "FAQ":
        st.text(''' 
        
        Frequently Asked Questions
        Let’s clear your Doubts!


    • Can I upload any format of files to convert it to LaTeX?
        As we provide three types of conversion, the format of files for each 
        conversion differs. 
        For Equation to LaTeX it supports jpg, jpeg, png
        For Table to LaTeX it supports pdf
        For Text to LaTeX it supports pdf, text.

    • Why does the Equation to LaTeX conversion take so long?
        At times it depends on the complexity of the equation, size of the image, 
        internet speed. If the equation is too long then it will take some time to 
        process it so that the accuracy level is not affected.

    • Can I upload an Excel Sheet to convert the table?
        Not the excel sheet but excel to pdf converted file will work to convert 
        any type of table irrespective to the complexity of the table.

    • If the image of an Equation is a slight blur, can I get the LaTeX code?
        Yes, the image will be converted to its desired output irrespective of the 
        quality of image so blur factor does not affect the output.

    • Is your website free to use?
        Yes, it is completely free.

    • Is it safe to upload my files to your website?
        Yes, it is safe to upload your files to our website. We take the security 
        and privacy of our users very seriously and have implemented several measures 
        to ensure that your files are protected.
        ''')

    if selected == "About":
        st.text('''
        On this platform We are Specifiacaly Focused on to the help Researcher ,Where 
        they can get Latex Code of all equation by Uploading it through this website .
        Latex Convertor makes life easy for technlogy enthusiastic where we help you 
        to get the latex  code of any equation in just few click ,simply by uploading it.
        latex Convertor help you to get the latex code of any equation,tables and excel 
        to get the latex code.
        You can get latex code of all the content by just uploading it in pdf,word,excel 
        and table format or you can also get the latex code of respective input by just 
        clicking the picture through web camera and uploading it on our website.
        Latex Convertor provides you the access to download the output(latex code ) in pdf 
        format  
                ''')

    if selected == "Help":
        st.text('''
        STEPS FOR COVERSION
            1. Equation to LaTeX
            2. Open a web browser and go to our LaTeX converter website. 
            3. Click on “Choose File” to select your image of the equation from your PC 
            or Mobile. After choosing, click on “Upload”.
            4. Wait for the website to process the equation and generate the LaTeX code.
            5. Once the LaTeX code is generated, copy it from the website.
            6. Open a LaTeX editor or document and paste the LaTeX code into it.
            7. Compile the LaTeX document to see the equation in the desired format.
            8. Check for any errors or formatting issues and make necessary changes.
            9. Save the document with the equation in the desired format.

        2.  Table to LaTeX
            1. Copy the table you want to convert from the source document or website and 
            paste it into a PDF file. Make sure that the table is formatted correctly and 
            contains all the necessary information.
            2. Open a web browser and go to our LaTeX converter website.
            3. Click on “Choose File” to select the PDF file from your PC or Mobile. After 
            choosing, click on “Upload” to upload the PDF file of the table.
            4. Copy the output: Once the conversion is complete, the website will provide 
            you with the LaTeX code for the table. Copy this code to use it in your LaTeX 
            document.
            5. Insert the table: Open your LaTeX document and paste the copied code where 
            you want the table to appear. Make sure that you have included the necessary 
            packages in your document preamble to render the table correctly.
            6. Compile the document: Save your document and compile it using your preferred 
            LaTeX compiler. The table should now appear in your document.

        3.  Text to LaTeX
            1. Copy the text you want to convert and paste it into a PDF file.
            2. Open a web browser and go to our LaTeX converter website. 
            3. Click on “Choose File” to select your image from your PC or Mobile. After 
            choosing, click on “Upload” to upload the PDF file which contains the required 
            text to be converted.
            4. Wait for the website to process the text and generate the LaTeX code. This may 
            take a few seconds or longer depending on the length and complexity of the text.
            5. Review the LaTeX code that is generated and make any necessary adjustments or 
            edits.
            6. Copy the final LaTeX code and use it as needed in your document or project.
        ''')


        

if __name__ == "__main__":
    app()